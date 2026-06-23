"""Generate professional PI/law-enforcement-style investigation reports for cases.

Supports PDF (fpdf2), DOCX (python-docx), HTML, CSV, and XLSX (openpyxl).
All export functions accept an optional ``investigator`` parameter (default
``"Unknown"``) that is embedded in the report header; the existing call-sites
in app/routes/cases.py do not need to change.
"""

import io
import json
import hashlib
import datetime
from collections import Counter

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _pdf_safe(text: str) -> str:
    """Strip characters Helvetica/latin-1 cannot encode to prevent fpdf2 crashes."""
    return str(text).encode("latin-1", errors="replace").decode("latin-1")


def _safe_str(v) -> str:
    """Return a clean string from any value, never None."""
    if v is None:
        return ""
    return str(v)


def _ts(dt) -> str:
    """Format a datetime to a human-readable string."""
    if dt is None:
        return "Unknown"
    return dt.strftime("%Y-%m-%d %H:%M UTC")


def _evidence_hash(result_json: str) -> str:
    """Return a short SHA-256 fingerprint of the raw result JSON."""
    h = hashlib.sha256(result_json.encode("utf-8", errors="replace")).hexdigest()
    return f"sha256:{h[:16]}..."


def _confidence_label(confidence) -> str:
    if not confidence:
        return "UNVERIFIED"
    return str(confidence).upper()


def _generated_now() -> str:
    return datetime.datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")


# ---------------------------------------------------------------------------
# Image download helper (used by PDF, DOCX, and HTML exports)
# ---------------------------------------------------------------------------

def _fetch_image_bytes(url: str, max_bytes: int = 2_097_152):
    """Download image → (bytes, mime_type) or (None, None). Never raises."""
    if not url:
        return None, None
    try:
        import httpx
        with httpx.Client(timeout=3, follow_redirects=True) as client:
            r = client.get(url, headers={
                "User-Agent": "Mozilla/5.0 (compatible; OSINT-Tracker/1.0)",
                "Referer": url,
            })
        if r.status_code != 200:
            return None, None
        mime = r.headers.get("content-type", "image/jpeg").split(";")[0].strip()
        return r.content[:max_bytes], mime
    except Exception:
        return None, None


def _prefetch_images(investigations) -> dict:
    """Concurrently fetch all profile images. Returns {url: (bytes, mime)}."""
    from concurrent.futures import ThreadPoolExecutor, as_completed
    urls = set()
    for inv in investigations:
        if inv.kind != "social" or not inv.result_json:
            continue
        try:
            d = json.loads(inv.result_json)
        except Exception:
            continue
        for r in d.get("results", []):
            url = r.get("profile_image") or r.get("image", "")
            if url:
                urls.add(url)
    if not urls:
        return {}
    cache = {}
    with ThreadPoolExecutor(max_workers=12) as pool:
        future_to_url = {pool.submit(_fetch_image_bytes, url): url for url in urls}
        for future in as_completed(future_to_url, timeout=20):
            url = future_to_url[future]
            try:
                cache[url] = future.result()
            except Exception:
                cache[url] = (None, None)
    return cache


def _report_fingerprint(case, investigations) -> str:
    """Stable SHA-256 fingerprint of the source data for chain-of-custody."""
    parts = [f"case:{case.id}", f"title:{case.title}"]
    for inv in sorted(investigations, key=lambda i: i.id or 0):
        parts.append(f"{inv.id}:{inv.kind}:{inv.query}:{hashlib.sha256((inv.result_json or '').encode()).hexdigest()[:8]}")
    return hashlib.sha256("|".join(parts).encode()).hexdigest()


_IMG_FMT = {
    "image/jpeg": "JPEG", "image/jpg": "JPEG",
    "image/png": "PNG", "image/gif": "GIF",
    "image/webp": "WEBP", "image/bmp": "BMP",
}


def _img_data_uri(url: str) -> str:
    """Fetch image and return a base64 data URI, or empty string on failure."""
    import base64
    img_bytes, mime = _fetch_image_bytes(url)
    if not img_bytes:
        return ""
    b64 = base64.b64encode(img_bytes).decode("ascii")
    return f"data:{mime};base64,{b64}"


# ---------------------------------------------------------------------------
# Per-kind findings extractor
# ---------------------------------------------------------------------------

def _extract_findings(kind: str, result_json: str) -> list:
    """Return [(label, value), ...] for a given investigation kind.

    Falls back gracefully if the JSON is malformed or keys are missing.
    """
    try:
        d = json.loads(result_json)
    except Exception:
        return [("Raw data", result_json[:500])]

    pairs = []

    def g(obj, *keys, default="N/A"):
        """Safely navigate nested dicts/lists."""
        cur = obj
        for k in keys:
            if isinstance(cur, dict):
                cur = cur.get(k)
            elif isinstance(cur, list) and isinstance(k, int):
                cur = cur[k] if k < len(cur) else None
            else:
                return default
            if cur is None:
                return default
        return _safe_str(cur) if cur is not None else default

    # ------------------------------------------------------------------
    if kind == "ip":
        pairs.append(("IP Address", g(d, "ip", default=g(d, "query", default="N/A"))))
        geo = d.get("geo") or d
        pairs.append(("Country", g(geo, "country")))
        pairs.append(("Region", g(geo, "regionName", default=g(geo, "region"))))
        pairs.append(("City", g(geo, "city")))
        pairs.append(("Organization", g(geo, "org", default=g(geo, "as"))))
        pairs.append(("ASN", g(geo, "as", default=g(geo, "asn"))))
        # Coordinates
        lat = geo.get("lat") or (geo.get("ll") or [None])[0]
        lon = geo.get("lon") or (geo.get("ll") or [None, None])[1]
        if lat and lon:
            pairs.append(("Coordinates", f"{lat}, {lon}"))
        # VirusTotal
        vt = d.get("virustotal", {})
        vt_stats = g(vt, "data", "attributes", "last_analysis_stats", default=None)
        if isinstance(vt_stats, dict):
            malicious = vt_stats.get("malicious", 0)
            suspicious = vt_stats.get("suspicious", 0)
            pairs.append(("VirusTotal Malicious", str(malicious)))
            pairs.append(("VirusTotal Suspicious", str(suspicious)))
            threat = "High" if int(malicious) > 0 else "Low"
            pairs.append(("Threat Level", threat))
        elif isinstance(d.get("virustotal"), dict):
            # flat vt structure
            vt_flat = d["virustotal"]
            mal = vt_flat.get("malicious", vt_flat.get("last_analysis_stats", {}).get("malicious", 0))
            pairs.append(("VirusTotal Malicious", str(mal)))
            pairs.append(("Threat Level", "High" if int(mal) > 0 else "Low"))
        # Shodan
        shodan = d.get("shodan") or {}
        ports = shodan.get("open_ports", [])
        if ports:
            pairs.append(("Open Ports", ", ".join(str(p) for p in ports[:20])))
        services = shodan.get("detected_services", [])
        if services:
            pairs.append(("Detected Services", ", ".join(str(s) for s in services[:10])))

    # ------------------------------------------------------------------
    elif kind == "domain":
        pairs.append(("Domain", g(d, "domain", default=g(d, "ldhName", default=g(d, "handle")))))
        pairs.append(("Registrar", g(d, "registrar", default=g(d, "registrarName"))))
        # RDAP ldhName
        pairs.append(("Handle / ID", g(d, "ldhName", default=g(d, "handle"))))
        # Dates
        for date_key, label in [
            ("createdDate", "Created"), ("created", "Created"),
            ("updatedDate", "Updated"), ("updated", "Updated"),
            ("expiresDate", "Expiry"), ("expires", "Expiry"),
        ]:
            val = d.get(date_key)
            if val:
                pairs.append((label, _safe_str(val)))
                break
        # name servers
        ns = d.get("nameservers") or d.get("nameServers") or []
        if isinstance(ns, list) and ns:
            pairs.append(("Name Servers", ", ".join(str(n) for n in ns[:8])))
        elif isinstance(ns, dict):
            pairs.append(("Name Servers", ", ".join(ns.values())))
        # Registrant
        reg_email = d.get("registrantEmail") or d.get("registrant", {}).get("email") if isinstance(d.get("registrant"), dict) else None
        reg_org = d.get("registrantOrganization") or (d.get("registrant", {}).get("organization") if isinstance(d.get("registrant"), dict) else None)
        if reg_email:
            pairs.append(("Registrant Email", reg_email))
        if reg_org:
            pairs.append(("Registrant Org", reg_org))
        # Status
        status = d.get("status") or d.get("statuses")
        if isinstance(status, list):
            pairs.append(("Status", ", ".join(str(s) for s in status[:4])))
        elif status:
            pairs.append(("Status", _safe_str(status)))

    # ------------------------------------------------------------------
    elif kind == "subdomain":
        pairs.append(("Domain", g(d, "domain")))
        subdomains = d.get("subdomains") or d.get("results") or []
        pairs.append(("Subdomains Found", str(len(subdomains))))
        for item in subdomains[:20]:
            if isinstance(item, dict):
                host = item.get("hostname") or item.get("subdomain") or ""
                ip = item.get("ip") or item.get("ip_address") or ""
                pairs.append(("Subdomain", f"{host} -> {ip}" if ip else host))
            else:
                pairs.append(("Subdomain", _safe_str(item)))

    # ------------------------------------------------------------------
    elif kind == "email":
        pairs.append(("Email Address", g(d, "email", default=g(d, "address"))))
        breaches = d.get("breaches")
        if breaches is None:
            pairs.append(("Breaches", "None"))
        elif isinstance(breaches, list):
            pairs.append(("Breach Count", str(len(breaches))))
            names = []
            for b in breaches[:5]:
                if isinstance(b, dict):
                    n = b.get("Name") or b.get("title") or b.get("name") or ""
                    if n:
                        names.append(n)
                elif isinstance(b, str):
                    names.append(b)
            if names:
                pairs.append(("Breach Names", ", ".join(names)))
        else:
            pairs.append(("Breaches", _safe_str(breaches)))
        # Hunter.io
        verification = d.get("verification") or d.get("hunter") or {}
        if isinstance(verification, dict):
            score = verification.get("score") or verification.get("deliverability")
            email_type = verification.get("email_type") or verification.get("type")
            if score is not None:
                pairs.append(("Deliverability Score", _safe_str(score)))
            if email_type:
                pairs.append(("Email Type", _safe_str(email_type)))

    # ------------------------------------------------------------------
    elif kind == "email_header":
        pairs.append(("From", g(d, "from", default=g(d, "From"))))
        pairs.append(("To", g(d, "to", default=g(d, "To"))))
        pairs.append(("Subject", g(d, "subject", default=g(d, "Subject"))))
        pairs.append(("Date", g(d, "date", default=g(d, "Date"))))
        # Originating IP
        orig_ip = d.get("received_from") or d.get("x_originating_ip") or d.get("originating_ip")
        if orig_ip:
            pairs.append(("Originating IP", _safe_str(orig_ip)))
        # Auth results
        pairs.append(("SPF Result", g(d, "spf", default=g(d, "SPF"))))
        pairs.append(("DKIM Result", g(d, "dkim", default=g(d, "DKIM"))))
        pairs.append(("DMARC Result", g(d, "dmarc", default=g(d, "DMARC"))))
        # Hops
        received = d.get("received") or d.get("received_chain") or []
        if isinstance(received, list):
            pairs.append(("Hops Count", str(len(received))))
        elif received:
            pairs.append(("Received", _safe_str(received)[:200]))

    # ------------------------------------------------------------------
    elif kind == "social":
        username = d.get("username", "")
        found_count = int(d.get("found_count") or 0)
        confirmed_count = int(d.get("confirmed_count") or 0)
        possible_count = found_count - confirmed_count
        total = int(d.get("total_checked") or 0)
        pairs.append(("Username", username))
        pairs.append(("Platforms Confirmed", str(confirmed_count)))
        pairs.append(("Platforms Possible", str(possible_count)))
        pairs.append(("Total Checked", str(total)))
        confirmed_shown = 0
        possible_shown = 0
        for r in d.get("results", []):
            if not r.get("found"):
                continue
            if r.get("confidence") == "high" and confirmed_shown < 15:
                url = r.get("url", "")
                pairs.append((f"Confirmed: {r.get('site', '')}", url))
                confirmed_shown += 1
            elif r.get("confidence") != "high" and possible_shown < 10:
                url = r.get("url", "")
                pairs.append((f"Possible: {r.get('site', '')}", url))
                possible_shown += 1

    # ------------------------------------------------------------------
    elif kind == "company":
        name = d.get("name") or d.get("company") or d.get("query") or "N/A"
        pairs.append(("Company Name", _safe_str(name)))
        # Registries
        for registry_key, registry_label in [
            ("us_edgar", "US EDGAR"),
            ("uk", "UK Companies House"),
            ("nigeria", "Nigeria CAC"),
            ("canada", "Canada CBCA"),
        ]:
            reg_data = d.get(registry_key)
            if reg_data is None:
                pairs.append((registry_label, "Not found"))
            elif isinstance(reg_data, list):
                pairs.append((registry_label, f"{len(reg_data)} result(s)"))
            elif isinstance(reg_data, dict):
                count = reg_data.get("total") or reg_data.get("count") or (1 if reg_data else 0)
                pairs.append((registry_label, f"{count} result(s)"))
            else:
                pairs.append((registry_label, _safe_str(reg_data)[:100]))
        # DuckDuckGo
        info = d.get("info") or d.get("duckduckgo") or {}
        if isinstance(info, dict):
            for field in ("website", "phone", "email", "address", "founded"):
                val = info.get(field)
                if val:
                    pairs.append((f"DDG {field.title()}", _safe_str(val)))
        # Google dork links
        dorks = d.get("google_dorks") or d.get("dork_results") or []
        for link in dorks[:3]:
            if isinstance(link, dict):
                url = link.get("url") or link.get("link") or ""
                pairs.append(("Google Dork", url))
            elif isinstance(link, str):
                pairs.append(("Google Dork", link))

    # ------------------------------------------------------------------
    elif kind == "file_forensics":
        pairs.append(("File Name", g(d, "filename", default=g(d, "file_name"))))
        pairs.append(("File Type", g(d, "file_type", default=g(d, "type"))))
        pairs.append(("File Size", g(d, "file_size", default=g(d, "size"))))
        pairs.append(("MIME Type", g(d, "mime_type", default=g(d, "mime"))))
        # Hashes
        pairs.append(("MD5 Hash", g(d, "md5", default=g(d, "hashes", "md5"))))
        pairs.append(("SHA256 Hash", g(d, "sha256", default=g(d, "hashes", "sha256"))))
        # EXIF / metadata
        meta = d.get("metadata") or d.get("exif") or d
        pairs.append(("Camera Make", g(meta, "Make", default=g(meta, "camera_make"))))
        pairs.append(("Camera Model", g(meta, "Model", default=g(meta, "camera_model"))))
        pairs.append(("Software", g(meta, "Software")))
        pairs.append(("Date Taken", g(meta, "DateTimeOriginal", default=g(meta, "date_taken"))))
        gps_coords = meta.get("GPS_Coordinates") or meta.get("gps_coordinates")
        if gps_coords:
            pairs.append(("GPS Coordinates", _safe_str(gps_coords)))
        gps_loc = d.get("location") or meta.get("location")
        if gps_loc:
            pairs.append(("GPS Location", _safe_str(gps_loc)))
        # Image
        width = meta.get("ImageWidth") or meta.get("width") or d.get("width")
        height = meta.get("ImageLength") or meta.get("ImageHeight") or meta.get("height") or d.get("height")
        if width and height:
            pairs.append(("Dimensions", f"{width} x {height}"))
        # Audio
        duration = d.get("duration") or meta.get("duration")
        bitrate = d.get("bitrate") or meta.get("bitrate")
        if duration:
            pairs.append(("Duration", _safe_str(duration)))
        if bitrate:
            pairs.append(("Bitrate", _safe_str(bitrate)))
        # PDF
        pages = d.get("page_count") or meta.get("page_count") or meta.get("Pages")
        if pages:
            pairs.append(("PDF Pages", _safe_str(pages)))
        # DOCX
        author = meta.get("Author") or meta.get("author") or d.get("author")
        revision = meta.get("Revision") or meta.get("revision") or d.get("revision")
        if author:
            pairs.append(("Document Author", _safe_str(author)))
        if revision:
            pairs.append(("Revision", _safe_str(revision)))

    # ------------------------------------------------------------------
    elif kind == "phone":
        pairs.append(("Phone Number", g(d, "phone_number", default=g(d, "number", default=g(d, "phone")))))
        valid = d.get("valid")
        pairs.append(("Valid", "Yes" if valid else ("No" if valid is False else "Unknown")))
        pairs.append(("Country", g(d, "country", default=g(d, "country_name"))))
        pairs.append(("Country Code", g(d, "country_code", default=g(d, "calling_code"))))
        pairs.append(("Carrier", g(d, "carrier", default=g(d, "operator"))))
        pairs.append(("Line Type", g(d, "line_type", default=g(d, "type"))))

    # ------------------------------------------------------------------
    elif kind == "mac":
        pairs.append(("MAC Address", g(d, "mac", default=g(d, "address"))))
        pairs.append(("Vendor", g(d, "vendor", default=g(d, "manufacturer", default=g(d, "company")))))
        pairs.append(("OUI", g(d, "oui", default=g(d, "prefix"))))

    # ------------------------------------------------------------------
    elif kind == "vehicle":
        summary = d.get("summary") or {}
        results = d.get("results") or {}
        if isinstance(results, list):
            results = results[0] if results else {}
        combined = {**summary, **results}
        for label, key in [
            ("VIN", "vin"), ("Year", "year"), ("Make", "make"), ("Model", "model"),
            ("Body Type", "body_type"), ("Drive Type", "drive_type"),
            ("Plant Country", "plant_country"), ("Engine", "engine"),
        ]:
            val = combined.get(key) or d.get(key)
            pairs.append((label, _safe_str(val) if val else "N/A"))

    # ------------------------------------------------------------------
    elif kind == "crypto":
        pairs.append(("Address", g(d, "address")))
        pairs.append(("Network / Chain", g(d, "network", default=g(d, "chain", default=g(d, "blockchain")))))
        pairs.append(("Balance", g(d, "balance", default=g(d, "final_balance"))))
        pairs.append(("Transaction Count", g(d, "tx_count", default=g(d, "n_tx", default=g(d, "total_transactions")))))
        pairs.append(("First Seen", g(d, "first_seen", default=g(d, "first_tx"))))
        pairs.append(("Last Seen", g(d, "last_seen", default=g(d, "last_tx"))))

    # ------------------------------------------------------------------
    elif kind == "imei":
        pairs.append(("IMEI", g(d, "imei", default=g(d, "tac"))))
        pairs.append(("Brand", g(d, "brand", default=g(d, "manufacturer"))))
        pairs.append(("Model", g(d, "model", default=g(d, "model_name"))))
        pairs.append(("OS", g(d, "os", default=g(d, "operating_system"))))
        pairs.append(("Internal Storage", g(d, "internal_storage", default=g(d, "storage"))))
        pairs.append(("RAM", g(d, "ram", default=g(d, "memory"))))
        pairs.append(("Release Year", g(d, "release_year", default=g(d, "year"))))
        # Catch any other top-level keys we haven't covered
        covered = {"imei", "tac", "brand", "manufacturer", "model", "model_name",
                   "os", "operating_system", "internal_storage", "storage",
                   "ram", "memory", "release_year", "year"}
        for k, v in d.items():
            if k not in covered and isinstance(v, (str, int, float, bool)):
                pairs.append((k.replace("_", " ").title(), _safe_str(v)))

    # ------------------------------------------------------------------
    elif kind == "darkweb":
        pairs.append(("Query", g(d, "query")))
        results = d.get("results") or d.get("data") or []
        pairs.append(("Total Results", str(d.get("total") or len(results))))
        for item in results[:10]:
            if isinstance(item, dict):
                title = item.get("title") or item.get("name") or ""
                url = item.get("url") or item.get("onion") or item.get("link") or ""
                pairs.append(("Result", f"{title} | {url}" if title else url))
            else:
                pairs.append(("Result", _safe_str(item)[:200]))

    # ------------------------------------------------------------------
    elif kind == "person":
        pairs.append(("Full Name", g(d, "name", default=g(d, "full_name"))))
        usernames = d.get("likely_usernames") or d.get("usernames") or []
        if isinstance(usernames, list) and usernames:
            pairs.append(("Likely Usernames", ", ".join(str(u) for u in usernames[:8])))
        links = d.get("investigation_links") or d.get("links") or []
        for link in links[:10]:
            if isinstance(link, dict):
                label = link.get("label") or link.get("name") or "Link"
                url = link.get("url") or link.get("href") or ""
                pairs.append((f"Link: {label}", url))
            elif isinstance(link, str):
                pairs.append(("Link", link))

    # ------------------------------------------------------------------
    else:
        # Generic fallback: surface top-level scalar fields
        for k, v in d.items():
            if isinstance(v, (str, int, float, bool)):
                pairs.append((k.replace("_", " ").title(), _safe_str(v)[:300]))
            elif isinstance(v, list) and v:
                preview = ", ".join(str(x) for x in v[:5])
                pairs.append((k.replace("_", " ").title(), preview[:300]))

    # Remove pure N/A pairs (keep labelled empties that carry meaning)
    return [(lbl, val) for lbl, val in pairs if val and val != "N/A"]


# ---------------------------------------------------------------------------
# Social profile helper (used by XLSX sheet + HTML cards)
# ---------------------------------------------------------------------------

def _social_profile_rows(result_json: str) -> list:
    """Return a list of dicts for each found social profile."""
    try:
        d = json.loads(result_json)
    except Exception:
        return []
    rows = []
    for r in d.get("results", []):
        if not r.get("found"):
            continue
        rows.append({
            "platform": r.get("site", ""),
            "confidence": "Confirmed" if r.get("confidence") == "high" else "Possible",
            "username": d.get("username", ""),
            "display_name": r.get("display_name", ""),
            "bio": r.get("bio", ""),
            "profile_url": r.get("url", ""),
            "profile_image_url": r.get("profile_image") or r.get("image", ""),
            "http_status": r.get("status_code", ""),
        })
    return rows


# ---------------------------------------------------------------------------
# Executive summary builder
# ---------------------------------------------------------------------------

def _build_executive_summary(case, investigations):
    """Return a dict of summary statistics."""
    total = len(investigations)
    kind_counts = Counter(inv.kind for inv in investigations)
    conf_counts = Counter(_confidence_label(getattr(inv, "confidence", None)) for inv in investigations)
    dates = [inv.created_at for inv in investigations if inv.created_at]
    date_range = ""
    if dates:
        mn = min(dates).strftime("%Y-%m-%d")
        mx = max(dates).strftime("%Y-%m-%d")
        date_range = f"{mn} to {mx}" if mn != mx else mn
    return {
        "total": total,
        "kind_counts": dict(kind_counts.most_common()),
        "confirmed": conf_counts.get("CONFIRMED", 0),
        "possible": conf_counts.get("POSSIBLE", 0),
        "unverified": conf_counts.get("UNVERIFIED", 0),
        "date_range": date_range or "N/A",
    }


def _risk_notes(investigations):
    """Auto-generate risk assessment bullet points."""
    notes = []
    for inv in investigations:
        try:
            d = json.loads(inv.result_json)
        except Exception:
            continue
        if inv.kind == "ip":
            vt = d.get("virustotal", {})
            stats = None
            if isinstance(vt, dict):
                stats = (vt.get("data") or {}).get("attributes", {}).get("last_analysis_stats")
                if not stats:
                    stats = vt.get("last_analysis_stats")
            if isinstance(stats, dict) and int(stats.get("malicious", 0)) > 0:
                notes.append(f"IP {inv.query} flagged as malicious by VirusTotal vendors.")
        elif inv.kind == "email":
            breaches = d.get("breaches")
            if isinstance(breaches, list) and len(breaches) > 0:
                notes.append(f"Email {inv.query} found in {len(breaches)} data breach(es).")
        elif inv.kind == "darkweb":
            results = d.get("results") or d.get("data") or []
            if results:
                notes.append(f"Dark web query '{inv.query}' returned {len(results)} result(s).")
        elif inv.kind == "social":
            confirmed = int(d.get("confirmed_count") or 0)
            if confirmed > 5:
                notes.append(f"Username '{inv.query}' confirmed on {confirmed} platforms - broad digital footprint.")
    if not notes:
        notes.append("No automated high-risk indicators detected. Manual review of findings recommended.")
    return notes


# ===========================================================================
# PDF EXPORT
# ===========================================================================

def export_pdf(case, investigations, investigator: str = "Unknown") -> bytes:
    """Return PDF bytes for a case and its investigations."""
    try:
        from fpdf import FPDF
    except ImportError:
        raise RuntimeError("fpdf2 not installed. Run: pip install fpdf2")

    # Colours
    NAVY  = (30, 41, 59)
    GREEN = (22, 163, 74)
    AMBER = (217, 119, 6)
    GRAY  = (107, 114, 128)
    LIGHT = (248, 250, 252)
    WHITE = (255, 255, 255)
    DARK  = (17, 24, 39)

    pdf = FPDF()
    pdf.set_margins(15, 15, 15)
    pdf.set_auto_page_break(auto=True, margin=15)

    def w():
        return pdf.epw

    def divider(color=(200, 200, 200)):
        pdf.set_draw_color(*color)
        pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
        pdf.ln(3)

    def section_heading(text, size=12):
        pdf.set_font("Helvetica", "B", size)
        pdf.set_text_color(*NAVY)
        pdf.cell(w(), 8, _pdf_safe(text), new_x="LMARGIN", new_y="NEXT")
        pdf.set_draw_color(*NAVY)
        pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
        pdf.ln(3)

    def key_value_row(label, value, shade=False):
        if shade:
            pdf.set_fill_color(*LIGHT)
        else:
            pdf.set_fill_color(*WHITE)
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(*GRAY)
        lw = 55
        pdf.cell(lw, 6, _pdf_safe(str(label) + ":"), fill=True)
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(*DARK)
        remaining = w() - lw
        # Use multi_cell for long values
        x_save = pdf.get_x()
        y_save = pdf.get_y()
        pdf.multi_cell(remaining, 6, _pdf_safe(str(value)[:300]), fill=True, new_x="LMARGIN", new_y="NEXT")

    summary = _build_executive_summary(case, investigations)
    img_cache = _prefetch_images(investigations)
    fingerprint = _report_fingerprint(case, investigations)

    # ------------------------------------------------------------------ COVER
    pdf.add_page()

    # Big title bar
    pdf.set_fill_color(*NAVY)
    pdf.rect(0, 0, pdf.w, 40, "F")
    pdf.set_y(12)
    pdf.set_font("Helvetica", "B", 22)
    pdf.set_text_color(*WHITE)
    pdf.cell(w(), 10, "INVESTIGATION REPORT", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(200, 210, 220)
    pdf.cell(w(), 7, "CONFIDENTIAL -- AUTHORISED PERSONNEL ONLY", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(14)

    # Case title
    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(*NAVY)
    pdf.cell(w(), 10, _pdf_safe(case.title), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    divider()

    # Meta table
    meta_rows = [
        ("Case ID", f"#{case.id}"),
        ("Status", case.status.replace("_", " ").title()),
        ("Priority", case.priority.title()),
        ("Investigator", investigator),
        ("Date Opened", _ts(case.created_at)),
        ("Report Generated", _generated_now()),
    ]
    for i, (lbl, val) in enumerate(meta_rows):
        key_value_row(lbl, val, shade=(i % 2 == 0))
    pdf.ln(4)

    if case.description:
        pdf.set_font("Helvetica", "I", 9)
        pdf.set_text_color(*GRAY)
        pdf.multi_cell(w(), 5, _pdf_safe(f"Description: {case.description}"),
                       new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

    # --------------------------------------------------------- EXECUTIVE SUMMARY
    pdf.add_page()
    section_heading("EXECUTIVE SUMMARY", size=13)

    # Overview paragraph
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(*DARK)
    overview = (
        f"This report covers investigation case \"{_pdf_safe(case.title)}\" "
        f"(ID #{case.id}), assigned to investigator {_pdf_safe(investigator)}. "
        f"The case is currently {case.status.replace('_', ' ').title()} with "
        f"{case.priority.title()} priority. "
        f"A total of {summary['total']} investigation(s) were conducted "
        f"spanning the period {summary['date_range']}."
    )
    pdf.multi_cell(w(), 6, overview, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    # Stats table
    stats_rows = [
        ("Total Investigations", str(summary["total"])),
        ("Confirmed Findings", str(summary["confirmed"])),
        ("Possible Findings", str(summary["possible"])),
        ("Unverified Findings", str(summary["unverified"])),
        ("Date Range", summary["date_range"]),
    ]
    for i, (lbl, val) in enumerate(stats_rows):
        key_value_row(lbl, val, shade=(i % 2 == 0))
    pdf.ln(4)

    # Breakdown by type
    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(*NAVY)
    pdf.cell(w(), 6, "Investigation Type Breakdown:", new_x="LMARGIN", new_y="NEXT")
    for i, (knd, cnt) in enumerate(summary["kind_counts"].items()):
        key_value_row(knd.replace("_", " ").title(), str(cnt), shade=(i % 2 == 0))
    pdf.ln(4)

    # Risk assessment
    section_heading("RISK ASSESSMENT")
    for note in _risk_notes(investigations):
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(*DARK)
        pdf.set_x(pdf.l_margin + 4)
        pdf.multi_cell(w() - 4, 5, _pdf_safe(f"- {note}"), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    # Investigation timeline
    section_heading("INVESTIGATION TIMELINE")
    pdf.set_font("Helvetica", "B", 8)
    pdf.set_text_color(*WHITE)
    pdf.set_fill_color(*NAVY)
    col_w = [35, 28, 80, 28]
    for hdr, cw in zip(["Timestamp", "Type", "Query", "Confidence"], col_w):
        pdf.cell(cw, 6, hdr, fill=True)
    pdf.ln()
    for i, inv in enumerate(sorted(investigations, key=lambda x: x.created_at or datetime.datetime.min)):
        shade = (i % 2 == 0)
        if shade:
            pdf.set_fill_color(*LIGHT)
        else:
            pdf.set_fill_color(*WHITE)
        conf = _confidence_label(getattr(inv, "confidence", None))
        pdf.set_font("Helvetica", "", 8)
        pdf.set_text_color(*DARK)
        ts_str = inv.created_at.strftime("%Y-%m-%d %H:%M") if inv.created_at else ""
        pdf.cell(col_w[0], 5, _pdf_safe(ts_str), fill=shade)
        pdf.cell(col_w[1], 5, _pdf_safe(inv.kind.replace("_", " ").title()), fill=shade)
        query_short = _pdf_safe(inv.query[:38]) if inv.query else ""
        pdf.cell(col_w[2], 5, query_short, fill=shade)
        # Confidence badge colour
        if conf == "CONFIRMED":
            pdf.set_text_color(*GREEN)
        elif conf == "POSSIBLE":
            pdf.set_text_color(*AMBER)
        else:
            pdf.set_text_color(*GRAY)
        pdf.cell(col_w[3], 5, conf, fill=shade)
        pdf.set_text_color(*DARK)
        pdf.ln()
    pdf.ln(6)

    # --------------------------------------------------------- PER INVESTIGATION
    for inv in investigations:
        pdf.add_page()
        conf = _confidence_label(getattr(inv, "confidence", None))
        kind_label = inv.kind.replace("_", " ").upper()

        # Colour bar
        if conf == "CONFIRMED":
            bar_color = GREEN
        elif conf == "POSSIBLE":
            bar_color = AMBER
        else:
            bar_color = GRAY

        pdf.set_fill_color(*bar_color)
        pdf.rect(0, pdf.get_y(), pdf.w, 12, "F")
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(*WHITE)
        heading_text = f"[{conf}]  {kind_label}  --  {inv.query}"
        pdf.cell(w(), 12, _pdf_safe(heading_text[:90]), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(3)

        # Timestamp + evidence hash
        pdf.set_font("Helvetica", "", 8)
        pdf.set_text_color(*GRAY)
        ts_str = _ts(inv.created_at)
        ev_hash = _evidence_hash(inv.result_json)
        pdf.cell(w(), 5, _pdf_safe(f"Collected at: {ts_str}    |    {ev_hash}"),
                 new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)
        divider()

        # Findings
        findings = _extract_findings(inv.kind, inv.result_json)
        if findings:
            for i, (lbl, val) in enumerate(findings):
                key_value_row(lbl, val, shade=(i % 2 == 0))
        else:
            pdf.set_font("Helvetica", "I", 9)
            pdf.set_text_color(*GRAY)
            pdf.cell(w(), 6, "(No structured findings extracted)", new_x="LMARGIN", new_y="NEXT")

        # Social: append profile photo thumbnails after findings
        if inv.kind == "social":
            profiles = _social_profile_rows(inv.result_json)
            if profiles:
                pdf.ln(4)
                pdf.set_font("Helvetica", "B", 9)
                pdf.set_text_color(*NAVY)
                pdf.cell(w(), 5, "PROFILES FOUND", new_x="LMARGIN", new_y="NEXT")
                pdf.ln(2)
                for p in profiles:
                    img_bytes, mime = img_cache.get(p["profile_image_url"]) or _fetch_image_bytes(p["profile_image_url"])
                    thumb_size = 18  # mm
                    x0, y0 = pdf.get_x(), pdf.get_y()
                    if img_bytes:
                        fmt = _IMG_FMT.get(mime, "JPEG")
                        try:
                            pdf.image(io.BytesIO(img_bytes), x=x0, y=y0, h=thumb_size, type=fmt)
                        except Exception:
                            img_bytes = None
                    if not img_bytes:
                        pdf.set_fill_color(75, 85, 99)
                        pdf.rect(x0, y0, thumb_size, thumb_size, "F")
                        pdf.set_xy(x0 + 3, y0 + 5)
                        pdf.set_font("Helvetica", "B", 10)
                        pdf.set_text_color(*WHITE)
                        pdf.cell(thumb_size - 6, 8, _pdf_safe(p["platform"][:1].upper()))
                    tx = x0 + thumb_size + 3
                    conf_color = GREEN if p["confidence"] == "Confirmed" else AMBER
                    pdf.set_fill_color(*conf_color)
                    pdf.set_xy(tx, y0)
                    pdf.set_font("Helvetica", "B", 8)
                    pdf.set_text_color(*WHITE)
                    pdf.cell(22, 4, _pdf_safe(p["confidence"].upper()), fill=True)
                    pdf.set_xy(tx + 24, y0)
                    pdf.set_font("Helvetica", "B", 9)
                    pdf.set_text_color(*NAVY)
                    pdf.cell(w() - tx - 24, 4, _pdf_safe(p["platform"]))
                    row_y = y0 + 5
                    for text, size, color in [
                        (p["display_name"], 8, (30, 41, 59)),
                        (p["bio"][:100] if p["bio"] else "", 7, GRAY),
                        (p["profile_url"][:90], 7, (99, 102, 241)),
                    ]:
                        if text:
                            pdf.set_xy(tx, row_y)
                            pdf.set_font("Helvetica", "", size)
                            pdf.set_text_color(*color)
                            pdf.cell(w() - tx, 4, _pdf_safe(text))
                            row_y += 4
                    pdf.set_xy(x0, y0 + thumb_size + 3)
                    divider()

    # Integrity footer on last page
    pdf.set_font("Helvetica", "", 7)
    pdf.set_text_color(*GRAY)
    pdf.ln(6)
    pdf.cell(w(), 4, _pdf_safe(f"Report integrity fingerprint (SHA-256): {fingerprint}"),
             new_x="LMARGIN", new_y="NEXT")
    pdf.cell(w(), 4, _pdf_safe(f"Generated: {_generated_now()} | Investigations: {len(investigations)}"),
             new_x="LMARGIN", new_y="NEXT")

    return pdf.output()


# ===========================================================================
# HTML EXPORT
# ===========================================================================

def export_html(case, investigations, investigator: str = "Unknown") -> str:
    """Return a standalone HTML report with professional dark-accented styling."""
    import html as _html
    esc = lambda s: _html.escape(str(s) if s is not None else "")

    summary = _build_executive_summary(case, investigations)
    risk_notes = _risk_notes(investigations)
    img_cache = _prefetch_images(investigations)
    fingerprint = _report_fingerprint(case, investigations)

    # ---- Cover / header metadata table
    meta_rows_html = ""
    meta_pairs = [
        ("Case ID", f"#{case.id}"),
        ("Status", case.status.replace("_", " ").title()),
        ("Priority", case.priority.title()),
        ("Investigator", investigator),
        ("Date Opened", _ts(case.created_at)),
        ("Report Generated", _generated_now()),
    ]
    for lbl, val in meta_pairs:
        meta_rows_html += f"<tr><th>{esc(lbl)}</th><td>{esc(val)}</td></tr>\n"

    # ---- Executive summary
    type_breakdown_html = ""
    for knd, cnt in summary["kind_counts"].items():
        type_breakdown_html += f"<tr><td>{esc(knd.replace('_',' ').title())}</td><td>{cnt}</td></tr>\n"

    risk_html = "".join(f"<li>{esc(n)}</li>" for n in risk_notes)

    # ---- Timeline
    sorted_invs = sorted(investigations, key=lambda x: x.created_at or datetime.datetime.min)
    timeline_rows = ""
    for inv in sorted_invs:
        conf = _confidence_label(getattr(inv, "confidence", None))
        ts_str = esc(inv.created_at.strftime("%Y-%m-%d %H:%M") if inv.created_at else "")
        badge_class = "badge-confirmed" if conf == "CONFIRMED" else ("badge-possible" if conf == "POSSIBLE" else "badge-unverified")
        timeline_rows += (
            f"<tr>"
            f"<td class='ts-cell'>{ts_str}</td>"
            f"<td>{esc(inv.kind.replace('_',' ').title())}</td>"
            f"<td><code>{esc(inv.query)}</code></td>"
            f"<td><span class='badge {badge_class}'>{esc(conf)}</span></td>"
            f"</tr>\n"
        )

    # ---- Per-investigation cards
    inv_cards = ""
    for inv in investigations:
        conf = _confidence_label(getattr(inv, "confidence", None))
        kind_label = esc(inv.kind.replace("_", " ").title())
        ts_str = esc(_ts(inv.created_at))
        ev_hash = esc(_evidence_hash(inv.result_json))
        bar_class = "bar-confirmed" if conf == "CONFIRMED" else ("bar-possible" if conf == "POSSIBLE" else "bar-unverified")
        badge_class = "badge-confirmed" if conf == "CONFIRMED" else ("badge-possible" if conf == "POSSIBLE" else "badge-unverified")

        # Findings table
        findings = _extract_findings(inv.kind, inv.result_json)

        if inv.kind == "social":
            # Special: render as profile cards
            profiles = _social_profile_rows(inv.result_json)
            findings_html = ""
            for p in profiles:
                img_onerror = "this.style.display='none'"
                if p["profile_image_url"]:
                    cached = img_cache.get(p["profile_image_url"])
                    if cached and cached[0]:
                        import base64 as _b64
                        b64 = _b64.b64encode(cached[0]).decode("ascii")
                        data_uri = f"data:{cached[1]};base64,{b64}"
                    else:
                        data_uri = _img_data_uri(p["profile_image_url"])
                    src = data_uri if data_uri else esc(p["profile_image_url"])
                    img_tag = (
                        f'<img src="{src}" '
                        f'style="width:40px;height:40px;border-radius:50%;object-fit:cover;margin-right:10px;flex-shrink:0" '
                        f'onerror="{img_onerror}">'
                    )
                else:
                    img_tag = '<div class="avatar-placeholder"></div>'
                conf_cls = "badge-confirmed" if p["confidence"] == "Confirmed" else "badge-possible"
                bio_html = f'<p class="profile-bio">{esc(p["bio"][:200])}</p>' if p["bio"] else ""
                dn_html = f'<p class="profile-dn">{esc(p["display_name"])}</p>' if p["display_name"] else ""
                findings_html += f"""
<div class="profile-card">
  {img_tag}
  <div class="profile-info">
    <div class="profile-header">
      <strong>{esc(p["platform"])}</strong>
      <span class="badge {conf_cls}">{esc(p["confidence"])}</span>
    </div>
    {dn_html}
    {bio_html}
    <a href="{esc(p['profile_url'])}" class="profile-url" target="_blank" rel="noopener">{esc(p['profile_url'])}</a>
  </div>
</div>"""
            if not profiles:
                findings_html = "<p class='no-data'>No profiles found.</p>"
        else:
            if findings:
                rows_html = ""
                for i, (lbl, val) in enumerate(findings):
                    shade = ' class="shaded"' if i % 2 == 0 else ""
                    # Make URLs clickable
                    val_str = esc(str(val))
                    if str(val).startswith("http"):
                        val_str = f'<a href="{esc(str(val))}" target="_blank" rel="noopener">{esc(str(val)[:80])}</a>'
                    rows_html += f"<tr{shade}><th>{esc(str(lbl))}</th><td>{val_str}</td></tr>\n"
                findings_html = f'<table class="findings-table">{rows_html}</table>'
            else:
                findings_html = "<p class='no-data'>No structured findings extracted.</p>"

        inv_cards += f"""
<div class="inv-card">
  <div class="inv-bar {bar_class}">
    <span class="inv-kind">{kind_label}</span>
    <code class="inv-query">{esc(inv.query)}</code>
    <span class="badge {badge_class}">{esc(conf)}</span>
  </div>
  <div class="inv-meta">
    <span>Collected: {ts_str}</span>
    <span class="ev-hash">{ev_hash}</span>
  </div>
  <div class="inv-findings">{findings_html}</div>
</div>
"""

    desc_html = f'<p class="case-desc">{esc(case.description)}</p>' if case.description else ""

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Investigation Report — {esc(case.title)}</title>
<style>
/* ===== Reset & Base ===== */
*, *::before, *::after {{ box-sizing: border-box; margin: 0; padding: 0; }}
body {{
  font-family: 'Segoe UI', system-ui, -apple-system, sans-serif;
  font-size: 14px;
  color: #111827;
  background: #f8fafc;
  line-height: 1.5;
}}
a {{ color: #4f46e5; text-decoration: none; }}
a:hover {{ text-decoration: underline; }}
code {{ font-family: 'Courier New', monospace; font-size: 0.85em; }}

/* ===== Layout ===== */
.report-wrapper {{ max-width: 1000px; margin: 0 auto; padding: 2rem 1.5rem; }}

/* ===== Cover ===== */
.cover {{
  background: #1e293b;
  color: #f8fafc;
  border-radius: 8px;
  padding: 2.5rem 2rem 2rem;
  margin-bottom: 1.5rem;
}}
.cover-eyebrow {{
  font-size: 0.7rem;
  letter-spacing: 0.15em;
  color: #94a3b8;
  text-transform: uppercase;
  margin-bottom: 0.5rem;
}}
.cover-title {{
  font-size: 2rem;
  font-weight: 700;
  color: #fff;
  margin-bottom: 0.25rem;
}}
.cover-subtitle {{
  font-size: 1rem;
  color: #94a3b8;
  margin-bottom: 1.5rem;
}}
.case-desc {{ color: #cbd5e1; font-size: 0.875rem; margin-top: 1rem; }}
.meta-table {{ width: 100%; border-collapse: collapse; margin-top: 1rem; }}
.meta-table th {{
  text-align: left; padding: 6px 12px;
  color: #94a3b8; font-weight: 600; font-size: 0.8rem;
  width: 160px; white-space: nowrap;
}}
.meta-table td {{
  padding: 6px 12px; color: #f1f5f9; font-size: 0.875rem;
  border-bottom: 1px solid #334155;
}}

/* ===== Sections ===== */
.section {{ background: #fff; border-radius: 8px; border: 1px solid #e5e7eb; margin-bottom: 1.5rem; overflow: hidden; }}
.section-hdr {{
  background: #1e293b; color: #fff;
  padding: 0.75rem 1.25rem; font-size: 0.75rem;
  font-weight: 700; letter-spacing: 0.1em; text-transform: uppercase;
}}
.section-body {{ padding: 1.25rem; }}

/* ===== Executive Summary stats ===== */
.stats-grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(140px, 1fr)); gap: 1rem; margin-bottom: 1.25rem; }}
.stat-box {{
  background: #f8fafc; border: 1px solid #e2e8f0;
  border-radius: 6px; padding: 0.75rem 1rem; text-align: center;
}}
.stat-box .stat-val {{ font-size: 1.5rem; font-weight: 700; color: #1e293b; }}
.stat-box .stat-lbl {{ font-size: 0.75rem; color: #64748b; }}
.breakdown-table {{ width: 100%; border-collapse: collapse; }}
.breakdown-table th, .breakdown-table td {{ padding: 6px 12px; text-align: left; border-bottom: 1px solid #f1f5f9; font-size: 0.875rem; }}
.breakdown-table th {{ color: #64748b; font-weight: 600; }}

/* ===== Risk ===== */
.risk-list {{ padding-left: 1.25rem; }}
.risk-list li {{ margin-bottom: 0.35rem; font-size: 0.875rem; color: #374151; }}

/* ===== Timeline ===== */
.tl-table {{ width: 100%; border-collapse: collapse; }}
.tl-table th {{ background: #1e293b; color: #fff; padding: 8px 12px; font-size: 0.75rem; text-align: left; }}
.tl-table td {{ padding: 7px 12px; font-size: 0.8rem; border-bottom: 1px solid #f1f5f9; }}
.tl-table tr:nth-child(even) td {{ background: #f8fafc; }}
.ts-cell {{ white-space: nowrap; color: #6b7280; }}

/* ===== Investigation Cards ===== */
.inv-card {{ background: #fff; border: 1px solid #e5e7eb; border-radius: 8px; margin-bottom: 1.25rem; overflow: hidden; }}
.inv-bar {{
  display: flex; align-items: center; gap: 0.75rem;
  padding: 0.6rem 1.1rem; flex-wrap: wrap;
}}
.bar-confirmed {{ background: #166534; }}
.bar-possible   {{ background: #92400e; }}
.bar-unverified {{ background: #374151; }}
.inv-kind {{ font-size: 0.75rem; font-weight: 700; color: #fff; text-transform: uppercase; letter-spacing: 0.05em; }}
.inv-query {{ color: #d1fae5; font-size: 0.8rem; }}
.bar-possible .inv-query {{ color: #fef3c7; }}
.bar-unverified .inv-query {{ color: #d1d5db; }}
.inv-meta {{
  display: flex; justify-content: space-between; align-items: center;
  padding: 0.4rem 1.1rem; font-size: 0.72rem; color: #9ca3af;
  background: #f9fafb; border-bottom: 1px solid #f3f4f6;
}}
.ev-hash {{ font-family: monospace; font-size: 0.7rem; }}
.inv-findings {{ padding: 1rem 1.1rem; }}

/* ===== Findings Table ===== */
.findings-table {{ width: 100%; border-collapse: collapse; }}
.findings-table th {{
  text-align: left; padding: 5px 10px;
  color: #6b7280; font-weight: 600; font-size: 0.78rem;
  width: 200px; white-space: nowrap; vertical-align: top;
}}
.findings-table td {{
  padding: 5px 10px; color: #111827; font-size: 0.82rem;
  word-break: break-word;
}}
.findings-table tr.shaded td, .findings-table tr.shaded th {{ background: #f8fafc; }}

/* ===== Social Profiles ===== */
.profile-card {{
  display: flex; align-items: flex-start;
  padding: 10px 0; border-bottom: 1px solid #f3f4f6;
}}
.profile-card:last-child {{ border-bottom: none; }}
.profile-card img {{ width: 40px; height: 40px; border-radius: 50%; object-fit: cover; margin-right: 10px; flex-shrink: 0; }}
.avatar-placeholder {{
  width: 40px; height: 40px; border-radius: 50%;
  background: #e5e7eb; margin-right: 10px; flex-shrink: 0;
}}
.profile-info {{ flex: 1; min-width: 0; }}
.profile-header {{ display: flex; align-items: center; gap: 8px; margin-bottom: 2px; }}
.profile-dn {{ font-size: 0.82rem; color: #374151; margin-bottom: 2px; }}
.profile-bio {{ font-size: 0.78rem; color: #6b7280; margin-bottom: 4px; }}
.profile-url {{ font-size: 0.78rem; color: #4f46e5; word-break: break-all; }}

/* ===== Badges ===== */
.badge {{
  display: inline-block; font-size: 0.65rem; font-weight: 700;
  padding: 2px 7px; border-radius: 3px; letter-spacing: 0.05em;
  white-space: nowrap; text-transform: uppercase;
}}
.badge-confirmed  {{ background: #dcfce7; color: #166534; }}
.badge-possible   {{ background: #fef3c7; color: #92400e; }}
.badge-unverified {{ background: #f3f4f6; color: #6b7280; }}

.no-data {{ color: #9ca3af; font-size: 0.85rem; font-style: italic; }}

/* ===== Print ===== */
@media print {{
  body {{ background: #fff; font-size: 12px; }}
  .cover {{ background: #1e293b !important; -webkit-print-color-adjust: exact; print-color-adjust: exact; }}
  .bar-confirmed, .bar-possible, .bar-unverified {{ -webkit-print-color-adjust: exact; print-color-adjust: exact; }}
  .inv-card {{ page-break-inside: avoid; }}
  .report-wrapper {{ max-width: 100%; padding: 0; }}
}}
</style>
</head>
<body>
<div class="report-wrapper">

<!-- COVER -->
<div class="cover">
  <p class="cover-eyebrow">Ethical OSINT Tracker &mdash; Confidential</p>
  <h1 class="cover-title">Investigation Report</h1>
  <p class="cover-subtitle">{esc(case.title)}</p>
  {desc_html}
  <table class="meta-table">
    {meta_rows_html}
  </table>
</div>

<!-- EXECUTIVE SUMMARY -->
<div class="section">
  <div class="section-hdr">Executive Summary</div>
  <div class="section-body">
    <div class="stats-grid">
      <div class="stat-box"><div class="stat-val">{summary['total']}</div><div class="stat-lbl">Total Investigations</div></div>
      <div class="stat-box"><div class="stat-val" style="color:#166534">{summary['confirmed']}</div><div class="stat-lbl">Confirmed</div></div>
      <div class="stat-box"><div class="stat-val" style="color:#92400e">{summary['possible']}</div><div class="stat-lbl">Possible</div></div>
      <div class="stat-box"><div class="stat-val" style="color:#6b7280">{summary['unverified']}</div><div class="stat-lbl">Unverified</div></div>
    </div>
    <p style="font-size:.875rem;color:#374151;margin-bottom:1rem">
      Activity period: <strong>{esc(summary['date_range'])}</strong>
    </p>
    <table class="breakdown-table">
      <thead><tr><th>Investigation Type</th><th>Count</th></tr></thead>
      <tbody>{type_breakdown_html}</tbody>
    </table>
  </div>
</div>

<!-- RISK ASSESSMENT -->
<div class="section">
  <div class="section-hdr">Risk Assessment</div>
  <div class="section-body">
    <ul class="risk-list">{risk_html}</ul>
  </div>
</div>

<!-- TIMELINE -->
<div class="section">
  <div class="section-hdr">Investigation Timeline</div>
  <div class="section-body" style="padding:0">
    <table class="tl-table">
      <thead><tr><th>Timestamp</th><th>Type</th><th>Query</th><th>Confidence</th></tr></thead>
      <tbody>{timeline_rows}</tbody>
    </table>
  </div>
</div>

<!-- INVESTIGATIONS -->
<div class="section-hdr" style="border-radius:8px 8px 0 0;margin-bottom:0">Per-Investigation Findings</div>
{inv_cards}

<div style="margin-top:2rem;padding:1rem;background:#f1f5f9;border-radius:6px;font-size:0.7rem;color:#64748b;font-family:monospace">
  <strong>Report integrity fingerprint (SHA-256):</strong> {fingerprint}<br>
  Generated: {_generated_now()} | Investigations: {len(investigations)}
</div>

</div><!-- /report-wrapper -->
</body>
</html>"""


# ===========================================================================
# DOCX EXPORT
# ===========================================================================

def export_docx(case, investigations, investigator: str = "Unknown") -> bytes:
    """Return DOCX bytes with professional PI report structure."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    img_cache = _prefetch_images(investigations)
    fingerprint = _report_fingerprint(case, investigations)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    def add_heading(text, level=1, color_rgb=(30, 41, 59)):
        h = doc.add_heading(text, level=level)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color_rgb)
        return h

    def add_kv(label, value, doc_ref=None):
        target = doc_ref or doc
        p = target.add_paragraph()
        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        r1.font.color.rgb = RGBColor(107, 114, 128)
        r1.font.size = Pt(10)
        r2 = p.add_run(str(value))
        r2.font.size = Pt(10)
        return p

    def add_table_row(table, label, value, shade=False):
        row = table.add_row()
        row.cells[0].text = str(label)
        row.cells[1].text = str(value)[:500]
        # Style label cell
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(107, 114, 128)
        for run in row.cells[1].paragraphs[0].runs:
            run.font.size = Pt(9)
        if shade:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "F8FAFC")
                tcPr.append(shd)

    # ---------------------------------------------------------------- Cover
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_para.add_run("INVESTIGATION REPORT")
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(30, 41, 59)

    sub_para = doc.add_paragraph()
    sub_run = sub_para.add_run("CONFIDENTIAL — AUTHORISED PERSONNEL ONLY")
    sub_run.font.size = Pt(9)
    sub_run.font.color.rgb = RGBColor(107, 114, 128)
    sub_run.italic = True

    doc.add_paragraph()
    add_heading(case.title, level=1)

    # Meta table
    meta_table = doc.add_table(rows=1, cols=2)
    meta_table.style = "Table Grid"
    hdr_cells = meta_table.rows[0].cells
    hdr_cells[0].text = "Field"
    hdr_cells[1].text = "Value"
    for cell in hdr_cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1E293B")
        tcPr.append(shd)

    meta_rows = [
        ("Case ID", f"#{case.id}"),
        ("Status", case.status.replace("_", " ").title()),
        ("Priority", case.priority.title()),
        ("Investigator", investigator),
        ("Date Opened", _ts(case.created_at)),
        ("Report Generated", _generated_now()),
    ]
    for i, (lbl, val) in enumerate(meta_rows):
        add_table_row(meta_table, lbl, val, shade=(i % 2 == 0))

    if case.description:
        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run(f"Description: {case.description}")
        r.italic = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(107, 114, 128)

    doc.add_page_break()

    # -------------------------------------------------------- Executive Summary
    summary = _build_executive_summary(case, investigations)
    add_heading("EXECUTIVE SUMMARY", level=1)

    overview = (
        f'This report covers case "{case.title}" (ID #{case.id}), '
        f"investigator: {investigator}. "
        f"Status: {case.status.replace('_', ' ').title()}, Priority: {case.priority.title()}. "
        f"{summary['total']} investigation(s) conducted over the period {summary['date_range']}."
    )
    doc.add_paragraph(overview)

    doc.add_paragraph()
    add_heading("Statistics", level=2)
    stats_table = doc.add_table(rows=1, cols=2)
    stats_table.style = "Table Grid"
    for i, (lbl, val) in enumerate([
        ("Total Investigations", str(summary["total"])),
        ("Confirmed Findings", str(summary["confirmed"])),
        ("Possible Findings", str(summary["possible"])),
        ("Unverified Findings", str(summary["unverified"])),
        ("Date Range", summary["date_range"]),
    ]):
        add_table_row(stats_table, lbl, val, shade=(i % 2 == 0))

    doc.add_paragraph()
    add_heading("Investigation Type Breakdown", level=2)
    for knd, cnt in summary["kind_counts"].items():
        add_kv(knd.replace("_", " ").title(), str(cnt))

    # -------------------------------------------------------- Risk Assessment
    doc.add_paragraph()
    add_heading("RISK ASSESSMENT", level=1)
    for note in _risk_notes(investigations):
        p = doc.add_paragraph(note, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(10)

    # -------------------------------------------------------- Timeline
    doc.add_page_break()
    add_heading("INVESTIGATION TIMELINE", level=1)
    tl_table = doc.add_table(rows=1, cols=4)
    tl_table.style = "Table Grid"
    for cell, hdr_text in zip(tl_table.rows[0].cells, ["Timestamp", "Type", "Query", "Confidence"]):
        cell.text = hdr_text
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)

    sorted_invs = sorted(investigations, key=lambda x: x.created_at or datetime.datetime.min)
    for inv in sorted_invs:
        conf = _confidence_label(getattr(inv, "confidence", None))
        row = tl_table.add_row()
        row.cells[0].text = inv.created_at.strftime("%Y-%m-%d %H:%M") if inv.created_at else ""
        row.cells[1].text = inv.kind.replace("_", " ").title()
        row.cells[2].text = inv.query[:60]
        row.cells[3].text = conf
        for cell in row.cells:
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(8)

    # -------------------------------------------------------- Per Investigation
    doc.add_page_break()
    add_heading("PER-INVESTIGATION FINDINGS", level=1)

    for inv in investigations:
        conf = _confidence_label(getattr(inv, "confidence", None))
        kind_label = inv.kind.replace("_", " ").upper()

        doc.add_paragraph()
        h2 = doc.add_heading(f"[{conf}] {kind_label} — {inv.query}", level=2)
        for run in h2.runs:
            if conf == "CONFIRMED":
                run.font.color.rgb = RGBColor(22, 101, 52)
            elif conf == "POSSIBLE":
                run.font.color.rgb = RGBColor(146, 64, 14)
            else:
                run.font.color.rgb = RGBColor(107, 114, 128)

        # Timestamp + hash
        meta_p = doc.add_paragraph()
        mr = meta_p.add_run(
            f"Collected: {_ts(inv.created_at)}    |    {_evidence_hash(inv.result_json)}"
        )
        mr.font.size = Pt(8)
        mr.font.color.rgb = RGBColor(107, 114, 128)
        mr.italic = True

        # Findings
        if inv.kind == "social":
            profiles = _social_profile_rows(inv.result_json)
            if profiles:
                add_heading("Profiles Found", level=3)
                for p in profiles:
                    tbl = doc.add_table(rows=1, cols=2)
                    tbl.style = "Table Grid"
                    img_cell = tbl.rows[0].cells[0]
                    text_cell = tbl.rows[0].cells[1]
                    img_cell.width = Cm(2.5)
                    img_bytes, _ = img_cache.get(p["profile_image_url"]) or _fetch_image_bytes(p["profile_image_url"])
                    if img_bytes:
                        try:
                            img_cell.paragraphs[0].add_run().add_picture(
                                io.BytesIO(img_bytes), width=Cm(2)
                            )
                        except Exception:
                            img_cell.text = p["platform"][:2]
                    else:
                        img_cell.text = p["platform"][:2]
                    tp = text_cell.paragraphs[0]
                    r1 = tp.add_run(f"{p['platform']}  ")
                    r1.bold = True
                    r1.font.size = Pt(10)
                    r2 = tp.add_run(f"[{p['confidence'].upper()}]")
                    r2.font.size = Pt(8)
                    if p["confidence"] == "Confirmed":
                        r2.font.color.rgb = RGBColor(22, 101, 52)
                    else:
                        r2.font.color.rgb = RGBColor(146, 64, 14)
                    for text, size, italic in [
                        (p["display_name"], 9, False),
                        (p["bio"][:200] if p["bio"] else "", 8, True),
                        (p["profile_url"], 8, False),
                    ]:
                        if text:
                            pp = text_cell.add_paragraph(text)
                            pp.runs[0].font.size = Pt(size)
                            pp.runs[0].italic = italic
                            if text == p["profile_url"]:
                                pp.runs[0].font.color.rgb = RGBColor(79, 70, 229)
                    doc.add_paragraph()
            else:
                p_no = doc.add_paragraph("No profiles found.")
                p_no.runs[0].italic = True
                p_no.runs[0].font.color.rgb = RGBColor(107, 114, 128)
        else:
            findings = _extract_findings(inv.kind, inv.result_json)
            if findings:
                findings_tbl = doc.add_table(rows=0, cols=2)
                findings_tbl.style = "Table Grid"
                for i, (lbl, val) in enumerate(findings):
                    add_table_row(findings_tbl, str(lbl), str(val)[:400], shade=(i % 2 == 0))
            else:
                p = doc.add_paragraph("No structured findings extracted.")
                p.runs[0].italic = True
                p.runs[0].font.color.rgb = RGBColor(107, 114, 128)

    doc.add_paragraph()
    fp_para = doc.add_paragraph()
    fp_run = fp_para.add_run(f"Integrity fingerprint (SHA-256): {fingerprint}\nGenerated: {_generated_now()}")
    fp_run.font.size = Pt(7)
    fp_run.font.color.rgb = RGBColor(107, 114, 128)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ===========================================================================
# CSV EXPORT
# ===========================================================================

def export_csv(case, investigations, investigator: str = "Unknown") -> bytes:
    """Return UTF-8 CSV bytes (BOM for Excel compatibility)."""
    import csv
    buf = io.StringIO()
    writer = csv.writer(buf)

    # Header block
    writer.writerow(["INVESTIGATION REPORT"])
    writer.writerow(["Case", case.title])
    writer.writerow(["Case ID", f"#{case.id}"])
    writer.writerow(["Status", case.status.replace("_", " ").title()])
    writer.writerow(["Priority", case.priority.title()])
    writer.writerow(["Investigator", investigator])
    writer.writerow(["Date Opened", _ts(case.created_at)])
    writer.writerow(["Description", case.description or ""])
    writer.writerow(["Generated", _generated_now()])
    writer.writerow([])

    # Column headers
    writer.writerow(["#", "Type", "Query", "Confidence", "Date", "Evidence Hash", "Summary"])

    for i, inv in enumerate(investigations, 1):
        conf = _confidence_label(getattr(inv, "confidence", None))
        ts_str = inv.created_at.strftime("%Y-%m-%d %H:%M") if inv.created_at else ""
        ev_hash = _evidence_hash(inv.result_json)
        findings = _extract_findings(inv.kind, inv.result_json)
        summary_parts = [f"{lbl}: {val}" for lbl, val in findings[:8]]
        summary = " | ".join(summary_parts)[:500]
        writer.writerow([
            i,
            inv.kind.replace("_", " ").title(),
            inv.query,
            conf,
            ts_str,
            ev_hash,
            summary,
        ])

    return b"\xef\xbb\xbf" + buf.getvalue().encode("utf-8")


# ===========================================================================
# XLSX EXPORT
# ===========================================================================

def export_xlsx(case, investigations, investigator: str = "Unknown") -> bytes:
    """Return XLSX bytes using openpyxl with per-kind formatted summaries."""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
    except ImportError:
        raise RuntimeError("openpyxl not installed. Run: pip install openpyxl")

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Report"

    # Styles
    title_font  = Font(bold=True, size=14, color="1E293B")
    hdr_font    = Font(bold=True, color="FFFFFF", size=9)
    hdr_fill    = PatternFill("solid", fgColor="1E293B")
    alt_fill    = PatternFill("solid", fgColor="F8FAFC")
    conf_fill   = PatternFill("solid", fgColor="DCFCE7")   # green
    poss_fill   = PatternFill("solid", fgColor="FEF3C7")   # amber
    unver_fill  = PatternFill("solid", fgColor="F3F4F6")   # gray
    center_align = Alignment(horizontal="center", wrap_text=True)
    wrap_align   = Alignment(wrap_text=True, vertical="top")

    # ---- Case metadata block
    ws.append(["INVESTIGATION REPORT"])
    ws["A1"].font = title_font
    ws.append(["Case", case.title])
    ws.append(["Case ID", f"#{case.id}"])
    ws.append(["Status", case.status.replace("_", " ").title()])
    ws.append(["Priority", case.priority.title()])
    ws.append(["Investigator", investigator])
    ws.append(["Date Opened", _ts(case.created_at)])
    if case.description:
        ws.append(["Description", case.description])
    ws.append(["Generated", _generated_now()])
    ws.append([])

    # ---- Column headers
    hdrs = ["#", "Type", "Query", "Confidence", "Date", "Evidence Hash", "Summary"]
    ws.append(hdrs)
    hdr_row = ws.max_row
    for col_i in range(1, len(hdrs) + 1):
        cell = ws.cell(row=hdr_row, column=col_i)
        cell.font = hdr_font
        cell.fill = hdr_fill
        cell.alignment = center_align

    # ---- Data rows
    for i, inv in enumerate(investigations, 1):
        conf = _confidence_label(getattr(inv, "confidence", None))
        ts_str = inv.created_at.strftime("%Y-%m-%d %H:%M") if inv.created_at else ""
        ev_hash = _evidence_hash(inv.result_json)
        findings = _extract_findings(inv.kind, inv.result_json)
        summary_parts = [f"{lbl}: {val}" for lbl, val in findings[:8]]
        summary_str = " | ".join(summary_parts)[:600]

        ws.append([
            i,
            inv.kind.replace("_", " ").title(),
            inv.query,
            conf,
            ts_str,
            ev_hash,
            summary_str,
        ])
        row_num = ws.max_row

        # Confidence colour
        if conf == "CONFIRMED":
            row_fill = conf_fill
        elif conf == "POSSIBLE":
            row_fill = poss_fill
        else:
            row_fill = unver_fill if i % 2 == 0 else alt_fill

        for col_i in range(1, len(hdrs) + 1):
            cell = ws.cell(row=row_num, column=col_i)
            if conf in ("CONFIRMED", "POSSIBLE"):
                cell.fill = row_fill
            elif i % 2 == 0:
                cell.fill = alt_fill
            cell.alignment = wrap_align

    # Column widths
    ws.column_dimensions["A"].width = 5
    ws.column_dimensions["B"].width = 18
    ws.column_dimensions["C"].width = 30
    ws.column_dimensions["D"].width = 14
    ws.column_dimensions["E"].width = 18
    ws.column_dimensions["F"].width = 22
    ws.column_dimensions["G"].width = 70

    # ---- Social Profiles sheet
    social_invs = [inv for inv in investigations if inv.kind == "social"]
    if social_invs:
        ws2 = wb.create_sheet("Social Profiles")
        sp_hdrs = ["Platform", "Confidence", "Username", "Display Name", "Bio", "Profile URL", "Image URL", "HTTP"]
        ws2.append(sp_hdrs)
        hdr2_row = ws2.max_row
        for col_i in range(1, len(sp_hdrs) + 1):
            cell = ws2.cell(row=hdr2_row, column=col_i)
            cell.font = hdr_font
            cell.fill = hdr_fill
            cell.alignment = center_align

        confirmed_fill = PatternFill("solid", fgColor="D1FAE5")
        possible_fill  = PatternFill("solid", fgColor="FEF3C7")

        for inv in social_invs:
            for row_data in _social_profile_rows(inv.result_json):
                ws2.append([
                    row_data["platform"], row_data["confidence"], row_data["username"],
                    row_data["display_name"], row_data["bio"],
                    row_data["profile_url"], row_data["profile_image_url"], row_data["http_status"],
                ])
                fill = confirmed_fill if row_data["confidence"] == "Confirmed" else possible_fill
                for col_i in range(1, len(sp_hdrs) + 1):
                    ws2.cell(row=ws2.max_row, column=col_i).fill = fill

        # Make profile URLs clickable
        for row_i in range(hdr2_row + 1, ws2.max_row + 1):
            url_cell = ws2.cell(row=row_i, column=6)
            if url_cell.value and str(url_cell.value).startswith("http"):
                url_cell.hyperlink = url_cell.value
                url_cell.font = Font(color="4338CA", underline="single")

        ws2.column_dimensions["A"].width = 20
        ws2.column_dimensions["B"].width = 12
        ws2.column_dimensions["C"].width = 16
        ws2.column_dimensions["D"].width = 25
        ws2.column_dimensions["E"].width = 50
        ws2.column_dimensions["F"].width = 50
        ws2.column_dimensions["G"].width = 50
        ws2.column_dimensions["H"].width = 8

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()
